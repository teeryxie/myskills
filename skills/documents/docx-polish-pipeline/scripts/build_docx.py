#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from docx.styles.style import WD_STYLE_TYPE
from docx.table import Table
from docx.text.paragraph import Paragraph


SKILL_DIR = Path(__file__).resolve().parents[1]
STYLE_CONFIG_PATH = SKILL_DIR / "assets" / "style-config.json"


def run(cmd: list[str], cwd: Path | None = None) -> None:
    subprocess.run(cmd, cwd=str(cwd) if cwd else None, check=True)


def load_profile(profile: str) -> dict:
    data = json.loads(STYLE_CONFIG_PATH.read_text(encoding="utf-8"))
    profiles = data["profiles"]
    selected = profile if profile in profiles else data["default_profile"]
    return profiles[selected]


def set_style_font(style, east_asia: str, latin: str, size_pt: float, bold: bool = False, color=None):
    font = style.font
    font.name = latin
    font.size = Pt(size_pt)
    font.bold = bold
    if color is not None:
        font.color.rgb = RGBColor(*color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def build_reference_docx(reference_docx: Path, profile_cfg: dict) -> None:
    reference_docx.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    margins = profile_cfg["page_margins_cm"]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.1)

    normal = doc.styles["Normal"]
    set_style_font(
        normal,
        east_asia=profile_cfg["body_font_east_asia"],
        latin=profile_cfg["body_font_latin"],
        size_pt=profile_cfg["body_size_pt"],
    )
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5)

    title = doc.styles["Title"]
    set_style_font(
        title,
        east_asia=profile_cfg["title_font_east_asia"],
        latin=profile_cfg["title_font_latin"],
        size_pt=profile_cfg["title_size_pt"],
        bold=True,
        color=profile_cfg["title_color_rgb"],
    )
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)

    if "Subtitle" not in doc.styles:
        subtitle = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = doc.styles["Subtitle"]
    set_style_font(
        subtitle,
        east_asia=profile_cfg["title_font_east_asia"],
        latin=profile_cfg["title_font_latin"],
        size_pt=profile_cfg["subtitle_size_pt"],
        color=(88, 102, 122),
    )
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    for style_name, size_key, color_key, before, after in [
        ("Heading 1", "heading1_size_pt", "heading1_color_rgb", 14, 7),
        ("Heading 2", "heading2_size_pt", "heading2_color_rgb", 12, 6),
        ("Heading 3", "heading3_size_pt", "heading3_color_rgb", 10, 5),
    ]:
        style = doc.styles[style_name]
        set_style_font(
            style,
            east_asia=profile_cfg["title_font_east_asia"],
            latin=profile_cfg["title_font_latin"],
            size_pt=profile_cfg[size_key],
            bold=True,
            color=profile_cfg[color_key],
        )
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Caption" not in doc.styles:
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption"]
    set_style_font(caption, east_asia="楷体", latin="Georgia", size_pt=10, color=(89, 89, 89))
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(reference_docx)


def render_mermaid_blocks(text: str, art_dir: Path) -> str:
    if not shutil.which("mmdc"):
        return text
    art_dir.mkdir(parents=True, exist_ok=True)
    pattern = re.compile(r"```mermaid\s*\n(.*?)\n```", re.DOTALL)
    counter = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal counter
        counter += 1
        mmd = art_dir / f"diagram-{counter}.mmd"
        png = art_dir / f"diagram-{counter}.png"
        mmd.write_text(match.group(1).strip() + "\n", encoding="utf-8")
        try:
            run(
                [
                    "mmdc",
                    "-i",
                    str(mmd),
                    "-o",
                    str(png),
                    "-b",
                    "white",
                    "-t",
                    "neutral",
                    "-w",
                    "1400",
                    "-s",
                    "2",
                    "-q",
                ]
            )
            return f"![diagram]({png.as_posix()})"
        except subprocess.CalledProcessError:
            return match.group(0)

    return pattern.sub(repl, text)


def preprocess_markdown(input_md: Path, processed_md: Path, art_dir: Path) -> None:
    text = input_md.read_text(encoding="utf-8")
    text = render_mermaid_blocks(text, art_dir)
    processed_md.write_text(text, encoding="utf-8")


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return "<w:drawing" in paragraph._p.xml


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width.twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_layout_fixed(table: Table) -> None:
    try:
        table.autofit = False
    except Exception:
        pass
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def choose_table_widths(headers: list[str], cols: int) -> list[float]:
    if headers == ["项目", "内容"]:
        return [4.0, 11.2]
    if headers == ["工作项", "人月数", "单价（元/人月）", "金额（元）", "主要职责"]:
        return [4.4, 1.6, 2.8, 2.3, 5.1]
    if headers == ["实施阶段", "周期", "主要工作", "金额（元）"]:
        return [4.2, 2.0, 7.0, 2.6]
    if cols == 2:
        return [4.6, 11.6]
    if cols == 3:
        return [3.4, 4.2, 8.7]
    if cols == 4:
        return [3.2, 3.0, 7.2, 2.8]
    return [16.2 / max(cols, 1)] * cols


def style_tables(doc: Document, profile_cfg: dict) -> None:
    for table in doc.tables:
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_layout_fixed(table)
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        widths = choose_table_widths(headers, len(table.columns))
        if table.rows:
            set_repeat_table_header(table.rows[0])
        header = True
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
                if c_idx < len(widths):
                    set_cell_width(cell, widths[c_idx])
                if header:
                    set_cell_shading(cell, profile_cfg["header_fill"])
                elif r_idx % 2 == 0:
                    set_cell_shading(cell, profile_cfg["stripe_fill"])
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.2
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(9.6 if len(table.columns) >= 4 else 10.2)
                        run.font.name = "宋体"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        if header:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(*profile_cfg["header_text_rgb"])
            header = False


def append_simple_field(paragraph: Paragraph, instruction: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def set_paragraph_border(paragraph: Paragraph, edge: str = "bottom", color: str = "D9E2F0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge_el = p_bdr.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        p_bdr.append(edge_el)
    edge_el.set(qn("w:val"), "single")
    edge_el.set(qn("w:sz"), "6")
    edge_el.set(qn("w:space"), "1")
    edge_el.set(qn("w:color"), color)


def style_headers_and_footers(doc: Document, header_text: str) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.text = ""
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_para.add_run(header_text)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(102, 102, 102)
        set_paragraph_border(header_para, edge="bottom", color="D6DEE8")

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.text = ""
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix = footer_para.add_run("第 ")
        prefix.font.name = "微软雅黑"
        prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        prefix.font.size = Pt(8.5)
        prefix.font.color.rgb = RGBColor(120, 120, 120)
        append_simple_field(footer_para, "PAGE")
        suffix = footer_para.add_run(" 页")
        suffix.font.name = "微软雅黑"
        suffix._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        suffix.font.size = Pt(8.5)
        suffix.font.color.rgb = RGBColor(120, 120, 120)


def style_paragraphs(doc: Document, profile_cfg: dict, title_text: str | None, subtitle_text: str | None) -> None:
    cover_done = False
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        if not cover_done and text and style_name in {"Heading 1", "Title", "Normal"}:
            para.style = doc.styles["Title"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(72)
            para.paragraph_format.space_after = Pt(10)
            if subtitle_text:
                subtitle = insert_paragraph_after(para, subtitle_text, "Subtitle")
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_done = True
            continue

        if style_name in {"Normal", "Body Text"}:
            if para._p.pPr is not None and para._p.pPr.numPr is not None:
                para.paragraph_format.line_spacing = 1.35
                para.paragraph_format.space_after = Pt(5)
                continue
            if text and len(text) > 25 and not paragraph_has_drawing(para):
                para.paragraph_format.first_line_indent = Cm(0.74)
            para.paragraph_format.line_spacing = 1.45
            para.paragraph_format.space_after = Pt(5.5)
        elif style_name == "List Paragraph":
            para.paragraph_format.line_spacing = 1.35
            para.paragraph_format.space_after = Pt(5)
        elif style_name == "Heading 1":
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
        elif style_name == "Heading 2":
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(5)

        for run in para.runs:
            if style_name in {"Normal", "Body Text", "List Paragraph"}:
                run.font.name = profile_cfg["body_font_latin"]
                run._element.rPr.rFonts.set(qn("w:eastAsia"), profile_cfg["body_font_east_asia"])


def style_images(doc: Document) -> None:
    max_width = Cm(15.5)
    max_height = Cm(9.8)
    for shape in doc.inline_shapes:
        width = float(shape.width)
        height = float(shape.height)
        scale = min(float(max_width) / width, float(max_height) / height)
        if scale > 0:
            shape.width = int(width * scale)
            shape.height = int(height * scale)


def detect_title(input_md: Path) -> str:
    for line in input_md.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return input_md.stem


def build_document(
    input_md: Path,
    output_docx: Path,
    profile: str = "formal-business",
    title: str | None = None,
    subtitle: str | None = None,
    build_dir: Path | None = None,
) -> Path:
    profile_cfg = load_profile(profile)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    actual_title = title or detect_title(input_md)
    actual_build_dir = build_dir or (output_docx.parent / f"{output_docx.stem}_build")
    actual_build_dir.mkdir(parents=True, exist_ok=True)
    processed_md = actual_build_dir / f"{output_docx.stem}-rendered.md"
    art_dir = actual_build_dir / "artifacts"
    reference_docx = actual_build_dir / "reference.docx"

    build_reference_docx(reference_docx, profile_cfg)
    preprocess_markdown(input_md, processed_md, art_dir)

    run(
        [
            "pandoc",
            str(processed_md),
            "--from",
            "markdown+pipe_tables",
            "--to",
            "docx",
            "--reference-doc",
            str(reference_docx),
            "--resource-path",
            f"{actual_build_dir}:{input_md.parent}",
            "-o",
            str(output_docx),
        ]
    )

    doc = Document(output_docx)
    style_paragraphs(doc, profile_cfg, actual_title, subtitle)
    style_tables(doc, profile_cfg)
    style_images(doc)
    style_headers_and_footers(doc, actual_title)
    doc.save(output_docx)
    return output_docx


def main() -> None:
    parser = argparse.ArgumentParser(description="Render markdown into a styled DOCX document.")
    parser.add_argument("--input-md", required=True, help="Markdown source file")
    parser.add_argument("--output-docx", required=True, help="Target DOCX path")
    parser.add_argument("--profile", default="formal-business", help="Style profile")
    parser.add_argument("--title", default="", help="Override title")
    parser.add_argument("--subtitle", default="", help="Optional subtitle")
    parser.add_argument("--build-dir", default="", help="Build directory for intermediates")
    args = parser.parse_args()

    output = build_document(
        input_md=Path(args.input_md).expanduser().resolve(),
        output_docx=Path(args.output_docx).expanduser().resolve(),
        profile=args.profile,
        title=args.title.strip() or None,
        subtitle=args.subtitle.strip() or None,
        build_dir=Path(args.build_dir).expanduser().resolve() if args.build_dir else None,
    )
    print(output)


if __name__ == "__main__":
    main()
